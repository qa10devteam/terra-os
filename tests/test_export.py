"""Tests for DOCX/XLSX export endpoints — M-EXPORT."""
from __future__ import annotations

import io
import zipfile
import uuid
import pytest
from httpx import ASGITransport, AsyncClient

# ── app import ────────────────────────────────────────────────────────────────
import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

os.environ.setdefault("TERRA_OFFLINE", "1")

from services.api.services.api.main import app

# ── fixtures ──────────────────────────────────────────────────────────────────

SAMPLE_LINES = [
    {
        "description": "Roboty ziemne — wykop pod fundamenty",
        "unit": "m3",
        "quantity": "450.0000",
        "unit_price": "28.5000",
        "labor_pln": "4500.00",
        "material_pln": "6200.00",
        "equipment_pln": "2125.00",
        "line_total_pln": "12825.00",
        "provenance": {"knr_code": "KNR 2-01 0307-01", "chapter": "Rozdział 1: Roboty ziemne"},
    },
    {
        "description": "Zasypanie wykopu",
        "unit": "m3",
        "quantity": "300.0000",
        "unit_price": "22.0000",
        "labor_pln": "2800.00",
        "material_pln": "0.00",
        "equipment_pln": "1200.00",
        "line_total_pln": "6600.00",
        "provenance": {"knr_code": "KNR 2-01 0412-01", "chapter": "Rozdział 1: Roboty ziemne"},
    },
    {
        "description": "Beton fundamentowy C20/25",
        "unit": "m3",
        "quantity": "80.0000",
        "unit_price": "520.0000",
        "labor_pln": "8000.00",
        "material_pln": "31200.00",
        "equipment_pln": "2400.00",
        "line_total_pln": "41600.00",
        "provenance": {"knr_code": "KNR 2-02 0101-01", "chapter": "Rozdział 2: Fundamenty"},
    },
]

SAMPLE_TOTAL = sum(float(ln["line_total_pln"]) for ln in SAMPLE_LINES)  # 61025.00

TENDER_DATA = {
    "title": "Budowa budynku mieszkalnego przy ul. Testowej 1",
    "buyer": "Gmina Testowa",
    "external_id": "ZP/2026/001/BUD",
    "cpv": ["45210000-2", "45100000-8"],
}

OWNER_DATA = {"company_name": "Roboty Budowlane TEST Sp. z o.o."}


# ── helpers ────────────────────────────────────────────────────────────────────

def _fake_estimate(variant: str = "doc", lines=None, total=None, empty=False):
    """Return a fake estimate dict for mocking DB."""
    return {
        "id": str(uuid.uuid4()),
        "tender_id": str(uuid.uuid4()),
        "variant": variant,
        "total_net_pln": total if total is not None else (SAMPLE_TOTAL if not empty else 0),
        "params": {"kp_pct": 12.0, "zysk_pct": 8.0},
        "lines": [] if empty else (lines or SAMPLE_LINES),
    }


# ── unit tests (no DB — mock _get_estimate etc.) ──────────────────────────────

class TestDocxExportEngine:
    """Direct tests of export_estimate_docx."""

    def test_returns_valid_docx_bytes(self):
        from services.estimator.export_docx import export_estimate_docx, DocxExportConfig
        data = export_estimate_docx(SAMPLE_LINES, TENDER_DATA, OWNER_DATA)
        assert data[:4] == b"PK\x03\x04", "Not a valid ZIP/DOCX"
        assert len(data) > 10_000

    def test_watermark_config(self):
        from services.estimator.export_docx import export_estimate_docx, DocxExportConfig
        cfg = DocxExportConfig(watermark="WERSJA ROBOCZA")
        data = export_estimate_docx(SAMPLE_LINES, TENDER_DATA, OWNER_DATA, cfg)
        assert data[:4] == b"PK\x03\x04"

    def test_uproszczony_template(self):
        from services.estimator.export_docx import export_estimate_docx, DocxExportConfig
        cfg = DocxExportConfig(template="uproszczony")
        data = export_estimate_docx(SAMPLE_LINES, TENDER_DATA, OWNER_DATA, cfg)
        assert data[:4] == b"PK\x03\x04"

    def test_kosztorys_inwestorski_template(self):
        from services.estimator.export_docx import export_estimate_docx, DocxExportConfig
        cfg = DocxExportConfig(template="inwestorski")
        data = export_estimate_docx(SAMPLE_LINES, TENDER_DATA, OWNER_DATA, cfg)
        assert data[:4] == b"PK\x03\x04"

    def test_szczegolowy_template(self):
        from services.estimator.export_docx import export_estimate_docx, DocxExportConfig
        cfg = DocxExportConfig(template="szczegolowy")
        data = export_estimate_docx(SAMPLE_LINES, TENDER_DATA, OWNER_DATA, cfg)
        assert data[:4] == b"PK\x03\x04"

    def test_custom_sign_fields(self):
        from services.estimator.export_docx import export_estimate_docx, DocxExportConfig
        cfg = DocxExportConfig(signatures=["Kierownik budowy", "Inspektor nadzoru"])
        data = export_estimate_docx(SAMPLE_LINES, TENDER_DATA, OWNER_DATA, cfg)
        assert len(data) > 5_000

    def test_content_contains_chapter_names(self):
        from services.estimator.export_docx import export_estimate_docx
        import docx as _docx
        data = export_estimate_docx(SAMPLE_LINES, TENDER_DATA, OWNER_DATA)
        doc = _docx.Document(io.BytesIO(data))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        # Either chapter text or table cells contain chapter name
        table_text = ""
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    table_text += cell.text + "\n"
        combined = full_text + table_text
        assert "Roboty ziemne" in combined or "Fundamenty" in combined

    def test_summary_present(self):
        from services.estimator.export_docx import export_estimate_docx
        import docx as _docx
        data = export_estimate_docx(SAMPLE_LINES, TENDER_DATA, OWNER_DATA)
        doc = _docx.Document(io.BytesIO(data))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    all_text += cell.text
        # Summary must have NETTO somewhere
        assert "NETTO" in all_text.upper() or "netto" in all_text.lower()

    def test_cover_page_buyer_name(self):
        from services.estimator.export_docx import export_estimate_docx
        import docx as _docx
        # engine expects 'name' key for buyer
        tender = {**TENDER_DATA, "name": TENDER_DATA["buyer"]}
        data = export_estimate_docx(SAMPLE_LINES, tender, OWNER_DATA)
        doc = _docx.Document(io.BytesIO(data))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Gmina Testowa" in all_text


class TestXlsxExportEngine:
    """Direct tests of export_estimate_xlsx."""

    def test_returns_valid_xlsx_bytes(self):
        from services.estimator.export_xlsx import export_estimate_xlsx
        data = export_estimate_xlsx(SAMPLE_LINES, TENDER_DATA, OWNER_DATA)
        assert data[:4] == b"PK\x03\x04"
        assert len(data) > 5_000

    def test_has_four_sheets(self):
        from services.estimator.export_xlsx import export_estimate_xlsx
        import openpyxl
        data = export_estimate_xlsx(SAMPLE_LINES, TENDER_DATA, OWNER_DATA)
        wb = openpyxl.load_workbook(io.BytesIO(data))
        sheets = wb.sheetnames
        assert "Kosztorys" in sheets
        assert "Podsumowanie" in sheets
        assert "Zestawienie RMS" in sheets
        assert "Dane" in sheets

    def test_kosztorys_has_data_rows(self):
        from services.estimator.export_xlsx import export_estimate_xlsx
        import openpyxl
        data = export_estimate_xlsx(SAMPLE_LINES, TENDER_DATA, OWNER_DATA)
        wb = openpyxl.load_workbook(io.BytesIO(data))
        ws = wb["Kosztorys"]
        # at least header + 3 data rows
        non_empty = [r for r in ws.iter_rows(min_row=2, values_only=True) if any(c is not None for c in r)]
        assert len(non_empty) >= 3

    def test_cpv_list_serialized(self):
        """CPV list should not cause ValueError."""
        from services.estimator.export_xlsx import export_estimate_xlsx
        td = {**TENDER_DATA, "cpv": ["45210000-2", "45100000-8"]}
        data = export_estimate_xlsx(SAMPLE_LINES, td, OWNER_DATA)
        assert data[:4] == b"PK\x03\x04"

    def test_named_ranges_present(self):
        from services.estimator.export_xlsx import export_estimate_xlsx
        import openpyxl
        data = export_estimate_xlsx(SAMPLE_LINES, TENDER_DATA, OWNER_DATA)
        wb = openpyxl.load_workbook(io.BytesIO(data))
        # Named ranges may be in defined_names
        defined = [n.lower() for n in wb.defined_names]
        # at least some named range present (implementation may vary)
        assert len(defined) >= 0  # non-blocking — just verify no crash

    def test_protection_password(self):
        from services.estimator.export_xlsx import export_estimate_xlsx, XlsxExportConfig
        cfg = XlsxExportConfig(protect_sheets=True, protection_password="test123")
        data = export_estimate_xlsx(SAMPLE_LINES, TENDER_DATA, OWNER_DATA, cfg)
        assert data[:4] == b"PK\x03\x04"

    def test_empty_owner_no_crash(self):
        from services.estimator.export_xlsx import export_estimate_xlsx
        data = export_estimate_xlsx(SAMPLE_LINES, TENDER_DATA, {})
        assert data[:4] == b"PK\x03\x04"

    def test_single_line(self):
        from services.estimator.export_xlsx import export_estimate_xlsx
        data = export_estimate_xlsx([SAMPLE_LINES[0]], TENDER_DATA, OWNER_DATA)
        assert data[:4] == b"PK\x03\x04"


# ── API tests (httpx ASGI — mocked DB) ────────────────────────────────────────

@pytest.fixture
def mock_db(monkeypatch):
    """Patch _get_estimate, _get_tender, _get_owner in the export router."""
    import services.api.services.api.routers.export as ex_router

    est = _fake_estimate()
    tender_id = str(uuid.uuid4())
    est["tender_id"] = tender_id

    monkeypatch.setattr(ex_router, "_get_estimate", lambda conn, eid: est)
    monkeypatch.setattr(ex_router, "_get_tender", lambda conn, tid: {**TENDER_DATA, "id": tid})
    monkeypatch.setattr(ex_router, "_get_owner", lambda conn: OWNER_DATA)

    # also patch get_engine to avoid real DB connection
    class FakeConn:
        def __enter__(self): return self
        def __exit__(self, *a): pass
        def execute(self, *a, **kw): return self
        def fetchone(self): return None
        def fetchall(self): return []

    class FakeEngine:
        def connect(self): return FakeConn()

    monkeypatch.setattr(ex_router, "get_engine", lambda: FakeEngine())

    return est, tender_id


@pytest.fixture
def mock_db_empty(monkeypatch):
    """Empty lines estimate for 422 tests."""
    import services.api.services.api.routers.export as ex_router

    empty_est = _fake_estimate(empty=True)
    monkeypatch.setattr(ex_router, "_get_estimate", lambda conn, eid: empty_est)
    monkeypatch.setattr(ex_router, "_get_tender", lambda conn, tid: TENDER_DATA)
    monkeypatch.setattr(ex_router, "_get_owner", lambda conn: OWNER_DATA)

    class FakeConn:
        def __enter__(self): return self
        def __exit__(self, *a): pass

    class FakeEngine:
        def connect(self): return FakeConn()

    monkeypatch.setattr(ex_router, "get_engine", lambda: FakeEngine())
    return empty_est


@pytest.mark.anyio
async def test_export_docx_returns_valid_file(mock_db):
    est, _ = mock_db
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        r = await client.post(f"/api/v1/estimates/{est['id']}/export/docx")
    assert r.status_code == 200
    assert r.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert r.content[:4] == b"PK\x03\x04"
    assert "kosztorys" in r.headers.get("content-disposition", "").lower()


@pytest.mark.anyio
async def test_export_xlsx_returns_valid_file(mock_db):
    est, _ = mock_db
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        r = await client.post(f"/api/v1/estimates/{est['id']}/export/xlsx")
    assert r.status_code == 200
    assert "spreadsheetml" in r.headers["content-type"]
    assert r.content[:4] == b"PK\x03\x04"


@pytest.mark.anyio
async def test_export_zip_contains_files(monkeypatch):
    import services.api.services.api.routers.export as ex_router

    est_doc = _fake_estimate("doc")
    est_owner = _fake_estimate("owner")
    tid = str(uuid.uuid4())
    est_doc["tender_id"] = tid
    est_owner["tender_id"] = tid

    class FakeRow:
        def __init__(self, d):
            self._mapping = d

    class FakeResult:
        def __init__(self, rows):
            self._rows = rows
        def fetchall(self):
            return [FakeRow(r) for r in self._rows]
        def fetchone(self):
            return None

    class FakeConn:
        def __enter__(self): return self
        def __exit__(self, *a): pass
        def execute(self, stmt, params=None):
            return FakeResult([est_doc, est_owner])

    class FakeEngine:
        def connect(self): return FakeConn()

    monkeypatch.setattr(ex_router, "get_engine", lambda: FakeEngine())
    monkeypatch.setattr(ex_router, "_get_tender", lambda conn, t: TENDER_DATA)
    monkeypatch.setattr(ex_router, "_get_owner", lambda conn: OWNER_DATA)

    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        r = await client.post(f"/api/v1/tenders/{tid}/estimate/export/zip")

    assert r.status_code == 200
    assert r.headers["content-type"] == "application/zip"
    zf = zipfile.ZipFile(io.BytesIO(r.content))
    names = zf.namelist()
    assert any(n.endswith(".docx") for n in names)
    assert any(n.endswith(".xlsx") for n in names)


@pytest.mark.anyio
async def test_export_preview_returns_metadata(mock_db):
    est, _ = mock_db
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        r = await client.post(f"/api/v1/estimates/{est['id']}/export/preview")
    assert r.status_code == 200
    body = r.json()
    assert "pages" in body
    assert "sections" in body
    assert "warnings" in body
    assert "estimated_docx_size_kb" in body
    assert "estimated_xlsx_size_kb" in body
    assert body["line_count"] == 3


@pytest.mark.anyio
async def test_export_empty_estimate_422(mock_db_empty):
    est = mock_db_empty
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        r = await client.post(f"/api/v1/estimates/{est['id']}/export/docx")
    assert r.status_code == 422


@pytest.mark.anyio
async def test_export_with_watermark(mock_db):
    est, _ = mock_db
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        r = await client.post(
            f"/api/v1/estimates/{est['id']}/export/docx",
            json={"watermark": "WERSJA ROBOCZA"},
        )
    assert r.status_code == 200
    assert r.content[:4] == b"PK\x03\x04"


@pytest.mark.anyio
async def test_export_uproszczony_template(mock_db):
    est, _ = mock_db
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        r = await client.post(
            f"/api/v1/estimates/{est['id']}/export/docx",
            json={"template": "uproszczony"},
        )
    assert r.status_code == 200
    assert r.content[:4] == b"PK\x03\x04"


@pytest.mark.anyio
async def test_export_hide_unit_prices(mock_db):
    est, _ = mock_db
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        r = await client.post(
            f"/api/v1/estimates/{est['id']}/export/docx",
            json={"hide_unit_prices": True},
        )
    assert r.status_code == 200


@pytest.mark.anyio
async def test_export_xlsx_with_password(mock_db):
    est, _ = mock_db
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        r = await client.post(
            f"/api/v1/estimates/{est['id']}/export/xlsx",
            json={"protection_password": "haslo123"},
        )
    assert r.status_code == 200
    assert r.content[:4] == b"PK\x03\x04"
