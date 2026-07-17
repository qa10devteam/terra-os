"""
Bud.OS Document Generator Service
Generates bid documents from Jinja2 templates → python-docx output.
Documents: Formularz Ofertowy, Kosztorys, Załączniki 1-4.
"""
import logging
from dataclasses import dataclass
from datetime import datetime
from enum import Enum
from io import BytesIO
from pathlib import Path
from typing import Any, Optional
from uuid import UUID, uuid4

from app.core.config import settings

logger = logging.getLogger(__name__)

TEMPLATES_DIR = Path(__file__).parent.parent / "templates" / "documents"


# ─── Models ─────────────────────────────────────────────────────────────────

class DocFormat(str, Enum):
    DOCX = "docx"
    PDF = "pdf"


@dataclass
class CompanyData:
    """Company profile data for document fill."""
    name: str
    nip: str
    regon: Optional[str] = None
    krs: Optional[str] = None
    address_street: str = ""
    address_city: str = ""
    address_zip: str = ""
    phone: str = ""
    email: str = ""
    website: str = ""
    reprezentant: str = ""  # person authorized to sign
    stanowisko: str = ""  # position/title


@dataclass
class TenderData:
    """Tender data for document context."""
    id: UUID
    bzp_number: Optional[str] = None
    title: str = ""
    zamawiajacy: str = ""
    zamawiajacy_address: str = ""
    deadline: Optional[datetime] = None
    procedure_type: str = ""
    wadium: Optional[float] = None


@dataclass
class BidData:
    """Bid data for document fill."""
    id: UUID
    price_net: float = 0.0
    price_vat: float = 0.0
    price_gross: float = 0.0
    termin_realizacji: str = ""  # e.g. "12 miesięcy"
    gwarancja: str = ""  # e.g. "60 miesięcy"
    termin_platnosci: str = "30 dni"


@dataclass
class GeneratedDoc:
    """Result of document generation."""
    id: UUID
    doc_type: str
    filename: str
    content: bytes
    file_size: int
    pages: int
    generated_at: datetime


# ─── Document Generator ────────────────────────────────────────────────────

class DocumentGenerator:
    """
    Generates bid documents using Jinja2 templates rendered to DOCX.
    
    Flow:
    1. Load Jinja2 template for document type
    2. Fill with context data (company, tender, bid, estimate)
    3. Render to python-docx Document
    4. Return bytes for S3 upload
    """

    def __init__(self):
        self._jinja_env = None

    @property
    def jinja_env(self):
        """Lazy-load Jinja2 environment."""
        if self._jinja_env is None:
            from jinja2 import Environment, FileSystemLoader
            self._jinja_env = Environment(
                loader=FileSystemLoader(str(TEMPLATES_DIR)),
                autoescape=False,
            )
        return self._jinja_env

    async def generate_formularz_ofertowy(
        self,
        company: CompanyData,
        tender: TenderData,
        bid: BidData,
    ) -> GeneratedDoc:
        """
        Generate Formularz Ofertowy (main bid form).
        Contains: company info, bid price, terms, declarations.
        """
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Header
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_para.add_run(f"{company.address_city}, {datetime.now().strftime('%d.%m.%Y')}")

        # Addressee
        doc.add_paragraph(f"\n{tender.zamawiajacy}")
        doc.add_paragraph(f"{tender.zamawiajacy_address}")

        # Title
        title = doc.add_heading("FORMULARZ OFERTOWY", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Reference
        doc.add_paragraph(
            f"Dotyczy: {tender.title}\n"
            f"Nr postępowania: {tender.bzp_number or 'N/A'}"
        )

        # Company data section
        doc.add_heading("1. DANE WYKONAWCY", level=2)
        doc.add_paragraph(f"Nazwa: {company.name}")
        doc.add_paragraph(f"Adres: {company.address_street}, {company.address_zip} {company.address_city}")
        doc.add_paragraph(f"NIP: {company.nip}")
        if company.regon:
            doc.add_paragraph(f"REGON: {company.regon}")
        if company.krs:
            doc.add_paragraph(f"KRS: {company.krs}")
        doc.add_paragraph(f"Tel: {company.phone}")
        doc.add_paragraph(f"Email: {company.email}")

        # Price section
        doc.add_heading("2. CENA OFERTY", level=2)
        doc.add_paragraph(
            f"Oferuję/oferujemy wykonanie przedmiotu zamówienia za cenę:\n\n"
            f"Cena netto: {bid.price_net:,.2f} PLN\n"
            f"VAT: {bid.price_vat:,.2f} PLN\n"
            f"Cena brutto: {bid.price_gross:,.2f} PLN\n"
            f"(słownie: {self._number_to_words(bid.price_gross)} złotych)"
        )

        # Terms
        doc.add_heading("3. WARUNKI REALIZACJI", level=2)
        doc.add_paragraph(f"Termin realizacji: {bid.termin_realizacji}")
        doc.add_paragraph(f"Okres gwarancji: {bid.gwarancja}")
        doc.add_paragraph(f"Termin płatności: {bid.termin_platnosci}")

        # Declarations
        doc.add_heading("4. OŚWIADCZENIA", level=2)
        declarations = [
            "Oświadczam/y, że zapoznałem/liśmy się z warunkami zamówienia i nie wnosimy zastrzeżeń.",
            "Oświadczam/y, że uzyskałem/liśmy wszystkie informacje niezbędne do przygotowania oferty.",
            "Oświadczam/y, że akceptujemy projekt umowy stanowiący załącznik do SWZ.",
            f"Oświadczam/y, że oferta jest ważna przez okres 30 dni od terminu składania ofert.",
        ]
        for i, decl in enumerate(declarations, 1):
            doc.add_paragraph(f"{i}. {decl}")

        # Signature
        doc.add_paragraph("\n\n")
        sign_para = doc.add_paragraph()
        sign_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sign_para.add_run("_________________________________\n")
        sign_para.add_run(f"{company.reprezentant}\n")
        sign_para.add_run(f"{company.stanowisko}")

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        content = buffer.getvalue()

        return GeneratedDoc(
            id=uuid4(),
            doc_type="formularz_ofertowy",
            filename=f"Formularz_Ofertowy_{tender.bzp_number or bid.id}.docx",
            content=content,
            file_size=len(content),
            pages=2,
            generated_at=datetime.utcnow(),
        )

    async def generate_kosztorys(
        self,
        company: CompanyData,
        tender: TenderData,
        bid: BidData,
        estimate_lines: list[dict],
    ) -> GeneratedDoc:
        """
        Generate Kosztorys Ofertowy (cost estimate document).
        Tabular format with all cost lines, subtotals, and grand total.
        """
        from docx import Document
        from docx.shared import Pt, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        title = doc.add_heading("KOSZTORYS OFERTOWY", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(
            f"Obiekt: {tender.title}\n"
            f"Zamawiający: {tender.zamawiajacy}\n"
            f"Wykonawca: {company.name}\n"
            f"Data: {datetime.now().strftime('%d.%m.%Y')}"
        )

        # Create table
        headers = ["Lp.", "Opis / KNR", "Jedn.", "Ilość", "R [PLN]", "M [PLN]", "S [PLN]", "Razem netto"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"

        # Header row
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        # Data rows
        for line in estimate_lines:
            row = table.add_row()
            row.cells[0].text = str(line.get("lp", ""))
            row.cells[1].text = f"{line.get('description', '')}\n{line.get('knr_code', '')}"
            row.cells[2].text = line.get("unit", "")
            row.cells[3].text = f"{line.get('quantity', 0):.2f}"
            row.cells[4].text = f"{line.get('r_total', 0):.2f}"
            row.cells[5].text = f"{line.get('m_total', 0):.2f}"
            row.cells[6].text = f"{line.get('s_total', 0):.2f}"
            row.cells[7].text = f"{line.get('net_total', 0):.2f}"

        # Summary
        doc.add_paragraph("\n")
        doc.add_heading("PODSUMOWANIE", level=2)
        doc.add_paragraph(f"Razem netto: {bid.price_net:,.2f} PLN")
        doc.add_paragraph(f"VAT: {bid.price_vat:,.2f} PLN")
        doc.add_paragraph(f"RAZEM BRUTTO: {bid.price_gross:,.2f} PLN")

        buffer = BytesIO()
        doc.save(buffer)
        content = buffer.getvalue()

        return GeneratedDoc(
            id=uuid4(),
            doc_type="kosztorys_ofertowy",
            filename=f"Kosztorys_Ofertowy_{tender.bzp_number or bid.id}.docx",
            content=content,
            file_size=len(content),
            pages=max(1, len(estimate_lines) // 30 + 1),
            generated_at=datetime.utcnow(),
        )

    async def generate_oswiadczenie_wykluczenie(
        self, company: CompanyData, tender: TenderData
    ) -> GeneratedDoc:
        """
        Generate Załącznik 1: Oświadczenie o braku podstaw wykluczenia.
        PZP art. 108 ust. 1, art. 109 ust. 1.
        """
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        title = doc.add_heading("OŚWIADCZENIE WYKONAWCY", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("o braku podstaw do wykluczenia z postępowania")

        doc.add_paragraph(
            f"\nDziałając w imieniu: {company.name}\n"
            f"NIP: {company.nip}\n"
            f"Adres: {company.address_street}, {company.address_zip} {company.address_city}\n"
        )

        doc.add_paragraph(
            "Oświadczam, że nie podlegam wykluczeniu z postępowania na podstawie:\n"
        )

        exclusion_grounds = [
            "art. 108 ust. 1 pkt 1-6 ustawy Pzp (obligatoryjne przesłanki wykluczenia)",
            "art. 109 ust. 1 pkt 1, 4-5, 7-10 ustawy Pzp (fakultatywne przesłanki wykluczenia, jeśli wskazane w SWZ)",
            "art. 7 ust. 1 ustawy z dnia 13 kwietnia 2022 r. o szczególnych rozwiązaniach (sankcje)",
        ]
        for ground in exclusion_grounds:
            doc.add_paragraph(f"• {ground}")

        doc.add_paragraph(
            f"\n\n{company.address_city}, dnia {datetime.now().strftime('%d.%m.%Y')}"
        )
        doc.add_paragraph(f"\n\n_________________________________\n{company.reprezentant}")

        buffer = BytesIO()
        doc.save(buffer)
        content = buffer.getvalue()

        return GeneratedDoc(
            id=uuid4(),
            doc_type="oswiadczenie_wykluczenie",
            filename=f"Zal_1_Oswiadczenie_Wykluczenie_{tender.bzp_number or 'draft'}.docx",
            content=content,
            file_size=len(content),
            pages=1,
            generated_at=datetime.utcnow(),
        )

    async def generate_wykaz_robot(
        self, company: CompanyData, tender: TenderData, roboty: list[dict]
    ) -> GeneratedDoc:
        """
        Generate Załącznik 2: Wykaz robót budowlanych.
        Lists completed projects as proof of experience.
        """
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        title = doc.add_heading("WYKAZ ROBÓT BUDOWLANYCH", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("(Załącznik nr 2 do SWZ)")

        doc.add_paragraph(f"\nWykonawca: {company.name}")

        headers = ["Lp.", "Rodzaj robót", "Wartość [PLN]", "Data wykonania", "Zamawiający"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        for i, robota in enumerate(roboty, 1):
            row = table.add_row()
            row.cells[0].text = str(i)
            row.cells[1].text = robota.get("opis", "")
            row.cells[2].text = f"{robota.get('wartosc', 0):,.2f}"
            row.cells[3].text = robota.get("data", "")
            row.cells[4].text = robota.get("zamawiajacy", "")

        doc.add_paragraph(f"\n\n{company.address_city}, dnia {datetime.now().strftime('%d.%m.%Y')}")
        doc.add_paragraph(f"\n_________________________________\n{company.reprezentant}")

        buffer = BytesIO()
        doc.save(buffer)
        content = buffer.getvalue()

        return GeneratedDoc(
            id=uuid4(),
            doc_type="wykaz_robot",
            filename=f"Zal_2_Wykaz_Robot_{tender.bzp_number or 'draft'}.docx",
            content=content,
            file_size=len(content),
            pages=1,
            generated_at=datetime.utcnow(),
        )

    async def generate_wykaz_osob(
        self, company: CompanyData, tender: TenderData, osoby: list[dict]
    ) -> GeneratedDoc:
        """
        Generate Załącznik 3: Wykaz osób (kadra techniczna).
        Lists key personnel with qualifications and construction permits.
        """
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        title = doc.add_heading("WYKAZ OSÓB", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("skierowanych do realizacji zamówienia (Załącznik nr 3 do SWZ)")

        doc.add_paragraph(f"\nWykonawca: {company.name}")

        headers = ["Lp.", "Imię i nazwisko", "Funkcja", "Uprawnienia", "Doświadczenie [lat]"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        for i, osoba in enumerate(osoby, 1):
            row = table.add_row()
            row.cells[0].text = str(i)
            row.cells[1].text = osoba.get("imie_nazwisko", "")
            row.cells[2].text = osoba.get("funkcja", "")
            row.cells[3].text = osoba.get("uprawnienia", "")
            row.cells[4].text = str(osoba.get("doswiadczenie_lat", ""))

        doc.add_paragraph(f"\n\n{company.address_city}, dnia {datetime.now().strftime('%d.%m.%Y')}")
        doc.add_paragraph(f"\n_________________________________\n{company.reprezentant}")

        buffer = BytesIO()
        doc.save(buffer)
        content = buffer.getvalue()

        return GeneratedDoc(
            id=uuid4(),
            doc_type="wykaz_osob",
            filename=f"Zal_3_Wykaz_Osob_{tender.bzp_number or 'draft'}.docx",
            content=content,
            file_size=len(content),
            pages=1,
            generated_at=datetime.utcnow(),
        )

    async def generate_all(
        self,
        company: CompanyData,
        tender: TenderData,
        bid: BidData,
        estimate_lines: list[dict],
        roboty: list[dict],
        osoby: list[dict],
    ) -> list[GeneratedDoc]:
        """Generate all bid documents."""
        docs = []
        docs.append(await self.generate_formularz_ofertowy(company, tender, bid))
        docs.append(await self.generate_kosztorys(company, tender, bid, estimate_lines))
        docs.append(await self.generate_oswiadczenie_wykluczenie(company, tender))
        docs.append(await self.generate_wykaz_robot(company, tender, roboty))
        docs.append(await self.generate_wykaz_osob(company, tender, osoby))
        return docs

    @staticmethod
    def _number_to_words(number: float) -> str:
        """Convert number to Polish words (simplified)."""
        # Simplified — in production use a proper library like `num2words`
        try:
            from num2words import num2words
            return num2words(int(number), lang='pl')
        except ImportError:
            return f"{number:,.2f}"
