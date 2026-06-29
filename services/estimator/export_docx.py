"""
Profesjonalny moduł eksportu kosztorysów budowlanych do DOCX.
Poziom komercyjnych programów kosztorysowych (Norma PRO, Zuzia, Rodos).

Autor: Terra.OS
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from datetime import date
from decimal import Decimal, ROUND_HALF_UP
from typing import Any, Dict, List, Literal, Optional, Tuple

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from num2words import num2words


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

@dataclass
class DocxExportConfig:
    """Pełna konfiguracja eksportu kosztorysu do DOCX."""

    # Typ szablonu
    template: Literal["ofertowy", "inwestorski", "uproszczony", "szczegolowy"] = "ofertowy"

    # Strona tytułowa
    title: str = "KOSZTORYS OFERTOWY"
    logo_path: Optional[str] = None  # ścieżka do logo (PNG/JPG)
    reference_number: str = ""
    cpv_codes: List[str] = field(default_factory=list)
    estimate_date: Optional[str] = None  # YYYY-MM-DD, default=today

    # Narzuty
    kp_percent: Decimal = Decimal("0")  # Koszty pośrednie %
    zysk_percent: Decimal = Decimal("0")  # Zysk %
    vat_percent: Decimal = Decimal("23")  # VAT %

    # Podpisy
    signatures: List[str] = field(default_factory=lambda: ["Sporządził", "Sprawdził", "Zatwierdził"])

    # Watermark
    watermark: Optional[str] = None  # np. "WERSJA ROBOCZA"

    # Klauzula waloryzacyjna
    valorization_clause: Optional[str] = None

    # Tabela elementów scalonych
    include_merged_elements: bool = False

    # Header / Footer
    custom_header: Optional[str] = None
    custom_footer: Optional[str] = None

    # Formatowanie
    currency: str = "PLN"
    page_orientation: Literal["portrait", "landscape"] = "portrait"

    # Marginesy (cm)
    margin_top: float = 2.0
    margin_bottom: float = 2.0
    margin_left: float = 2.5
    margin_right: float = 1.5


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _to_decimal(val: Any) -> Decimal:
    """Bezpiecznie konwertuj na Decimal."""
    if isinstance(val, Decimal):
        return val
    try:
        return Decimal(str(val).replace(",", ".").replace(" ", ""))
    except Exception:
        return Decimal("0")


def _format_pln(amount: Decimal) -> str:
    """Formatowanie kwoty w stylu PL: 12 825,00."""
    q = amount.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    sign = "-" if q < 0 else ""
    q = abs(q)
    int_part = int(q)
    dec_part = str(q).split(".")[-1] if "." in str(q) else "00"
    dec_part = dec_part.ljust(2, "0")[:2]
    # Separator tysięcy = spacja (non-breaking space)
    int_str = ""
    s = str(int_part)
    for i, ch in enumerate(reversed(s)):
        if i > 0 and i % 3 == 0:
            int_str = "\u00a0" + int_str
        int_str = ch + int_str
    return f"{sign}{int_str},{dec_part}"


def _amount_in_words(amount: Decimal) -> str:
    """Kwota słownie po polsku z walutą PLN."""
    q = amount.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    try:
        return num2words(float(q), lang="pl", to="currency", currency="PLN")
    except Exception:
        return num2words(float(q), lang="pl")


def _md_to_plain(text: Optional[str]) -> str:
    """Prosty markdown → plain text."""
    if not text:
        return ""
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"#{1,6}\s*", "", text)
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    return text.strip()


def _set_cell_border(cell, **kwargs):
    """Ustaw border na komórce tabeli. kwargs: top, bottom, left, right = dict(sz, color, val)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        + "".join(
            f'<w:{edge} w:val="{props.get("val", "single")}" '
            f'w:sz="{props.get("sz", "4")}" w:space="0" '
            f'w:color="{props.get("color", "000000")}"/>'
            for edge, props in kwargs.items()
        )
        + "</w:tcBorders>"
    )
    tcPr.append(tcBorders)


def _shade_cell(cell, color: str):
    """Ustaw kolor tła komórki."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _set_cell_text(cell, text: str, bold: bool = False, align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT, size: int = 9):
    """Ustaw tekst w komórce z formatowaniem."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    # Spacing
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)


def _apply_table_borders(table):
    """Dodaj ramki 0.5pt do całej tabeli."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def _add_page_number_footer(section):
    """Dodaj stopkę 'Strona X z Y'."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Strona ")
    run.font.size = Pt(8)
    # PAGE field
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run1 = p.add_run()
    run1._r.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2 = p.add_run()
    run2._r.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3 = p.add_run()
    run3._r.append(fldChar2)
    run4 = p.add_run(" z ")
    run4.font.size = Pt(8)
    # NUMPAGES field
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run5 = p.add_run()
    run5._r.append(fldChar3)
    instrText2 = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
    run6 = p.add_run()
    run6._r.append(instrText2)
    fldChar4 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run7 = p.add_run()
    run7._r.append(fldChar4)


def _add_watermark(doc: Document, text: str):
    """Dodaj watermark (znak wodny) do dokumentu."""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        # Użyj WordArt-style shape via VML
        r = p.add_run()
        # Add shape via fallback - simple diagonal text
        shape_xml = (
            f'<w:pict {nsdecls("w")}>'
            '<v:shapetype xmlns:v="urn:schemas-microsoft-com:vml" id="_x0000_t136" '
            'coordsize="21600,21600" adj="10800" path="m@7,l@8,m@5,21600l@6,21600e">'
            '</v:shapetype>'
            f'<v:shape xmlns:v="urn:schemas-microsoft-com:vml" '
            f'style="position:absolute;margin-left:0;margin-top:0;width:500pt;height:200pt;'
            f'rotation:315;z-index:-251657216;mso-position-horizontal:center;'
            f'mso-position-vertical:center" fillcolor="silver" stroked="f" type="#_x0000_t136">'
            f'<v:textpath style="font-family:Calibri;font-size:60pt" string="{text}"/>'
            '</v:shape></w:pict>'
        )
        try:
            r._r.append(parse_xml(shape_xml))
        except Exception:
            # Fallback: po prostu dodaj tekst w headerze z dużą czcionką
            p2 = header.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            wr = p2.add_run(text)
            wr.font.size = Pt(48)
            wr.font.color.rgb = RGBColor(0xC0, 0xC0, 0xC0)
            wr.bold = True


# ---------------------------------------------------------------------------
# Column definitions per template
# ---------------------------------------------------------------------------

COLUMNS = {
    "ofertowy": [
        ("Lp", 1.0, "center"),
        ("Podstawa", 2.5, "left"),
        ("Opis robót", 5.0, "left"),
        ("Jm", 1.0, "center"),
        ("Ilość", 1.5, "right"),
        ("Cena jdn.", 2.0, "right"),
        ("Robocizna", 2.0, "right"),
        ("Materiały", 2.0, "right"),
        ("Sprzęt", 2.0, "right"),
    ],
    "inwestorski": [
        ("Lp", 1.0, "center"),
        ("Podstawa", 2.5, "left"),
        ("Opis robót", 5.5, "left"),
        ("Jm", 1.0, "center"),
        ("Ilość", 1.5, "right"),
        ("Cena jdn.", 2.0, "right"),
        ("Robocizna", 2.0, "right"),
        ("Materiały", 2.0, "right"),
        ("Sprzęt", 2.0, "right"),
    ],
    "uproszczony": [
        ("Lp", 1.0, "center"),
        ("Podstawa", 2.5, "left"),
        ("Opis robót", 7.0, "left"),
        ("Jm", 1.0, "center"),
        ("Ilość", 2.0, "right"),
        ("Wartość", 2.5, "right"),
    ],
    "szczegolowy": [
        ("Lp", 0.8, "center"),
        ("Podstawa", 2.2, "left"),
        ("Opis robót", 4.5, "left"),
        ("Jm", 0.8, "center"),
        ("Ilość", 1.3, "right"),
        ("Cena jdn.", 1.5, "right"),
        ("Robocizna", 1.5, "right"),
        ("Materiały", 1.5, "right"),
        ("Sprzęt", 1.5, "right"),
        ("Wartość", 1.8, "right"),
    ],
}

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


# ---------------------------------------------------------------------------
# Main export function
# ---------------------------------------------------------------------------

def export_estimate_docx(
    lines: List[Dict[str, Any]],
    tender_data: Dict[str, Any],
    owner_data: Dict[str, Any],
    config: Optional[DocxExportConfig] = None,
) -> bytes:
    """
    Eksportuj kosztorys budowlany do DOCX.

    Args:
        lines: Lista pozycji kosztorysowych (dict z polami estimate line).
        tender_data: Dane zamawiającego (name, address, nip, phone, email).
        owner_data: Dane wykonawcy (name, address, nip, phone, email).
        config: Konfiguracja eksportu.

    Returns:
        Bytes dokumentu DOCX.
    """
    if config is None:
        config = DocxExportConfig()

    doc = Document()
    cols = COLUMNS[config.template]

    # --- Margins ---
    for section in doc.sections:
        section.top_margin = Cm(config.margin_top)
        section.bottom_margin = Cm(config.margin_bottom)
        section.left_margin = Cm(config.margin_left)
        section.right_margin = Cm(config.margin_right)
        if config.page_orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            new_width = section.page_height
            new_height = section.page_width
            section.page_width = new_width
            section.page_height = new_height

    # --- Watermark ---
    if config.watermark:
        _add_watermark(doc, config.watermark)

    # --- Custom header ---
    if config.custom_header:
        for section in doc.sections:
            header = section.header
            header.is_linked_to_previous = False
            hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            hp.text = _md_to_plain(config.custom_header)
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if hp.runs:
                hp.runs[0].font.size = Pt(8)

    # --- STRONA TYTUŁOWA ---
    _build_title_page(doc, lines, tender_data, owner_data, config)

    # --- Page break ---
    doc.add_page_break()

    # --- TABELA KOSZTORYSU ---
    _build_estimate_table(doc, lines, cols, config)

    # --- PODSUMOWANIE ---
    _build_summary(doc, lines, config)

    # --- Tabela elementów scalonych ---
    if config.include_merged_elements:
        _build_merged_elements_table(doc, lines, config)

    # --- Klauzula waloryzacyjna ---
    if config.valorization_clause:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run("Klauzula waloryzacyjna:")
        run.bold = True
        run.font.size = Pt(10)
        p2 = doc.add_paragraph()
        run2 = p2.add_run(config.valorization_clause)
        run2.font.size = Pt(9)
        run2.italic = True

    # --- PODPISY ---
    _build_signatures(doc, config)

    # --- Footer: Strona X z Y ---
    for section in doc.sections:
        _add_page_number_footer(section)
        # Custom footer text
        if config.custom_footer:
            footer = section.footer
            pf = footer.add_paragraph()
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            rf = pf.add_run(_md_to_plain(config.custom_footer))
            rf.font.size = Pt(7)

    # --- Serialize ---
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Title page
# ---------------------------------------------------------------------------

def _build_title_page(doc: Document, lines, tender_data, owner_data, config: DocxExportConfig):
    """Buduj stronę tytułową kosztorysu."""
    # Logo placeholder
    if config.logo_path:
        try:
            doc.add_picture(config.logo_path, width=Cm(4))
        except Exception:
            p = doc.add_paragraph("[LOGO]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("[LOGO]")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    # Tytuł
    title_map = {
        "ofertowy": "KOSZTORYS OFERTOWY",
        "inwestorski": "KOSZTORYS INWESTORSKI",
        "uproszczony": "KOSZTORYS UPROSZCZONY",
        "szczegolowy": "KOSZTORYS SZCZEGÓŁOWY",
    }
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run(config.title or title_map.get(config.template, "KOSZTORYS"))
    run_t.bold = True
    run_t.font.size = Pt(16)

    doc.add_paragraph()

    # Tabela informacyjna
    info_table = doc.add_table(rows=0, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    def _add_info_row(label: str, value: str):
        row = info_table.add_row()
        _set_cell_text(row.cells[0], label, bold=True, size=10)
        _set_cell_text(row.cells[1], value, size=10)

    est_date = config.estimate_date or date.today().strftime("%Y-%m-%d")
    _add_info_row("Data:", est_date)

    if config.reference_number:
        _add_info_row("Nr referencyjny:", config.reference_number)

    if config.cpv_codes:
        _add_info_row("Kody CPV:", ", ".join(config.cpv_codes))

    # Nazwa obiektu / inwestycji
    obj_name = tender_data.get("object_name", tender_data.get("name", ""))
    if obj_name:
        _add_info_row("Obiekt/Inwestycja:", obj_name)

    doc.add_paragraph()

    # Zamawiający
    p_z = doc.add_paragraph()
    run_z = p_z.add_run("ZAMAWIAJĄCY:")
    run_z.bold = True
    run_z.font.size = Pt(11)

    for k, label in [("name", ""), ("address", ""), ("nip", "NIP: "), ("phone", "Tel: "), ("email", "E-mail: ")]:
        val = tender_data.get(k, "")
        if val:
            pv = doc.add_paragraph()
            pv.paragraph_format.space_before = Pt(2)
            pv.paragraph_format.space_after = Pt(2)
            rv = pv.add_run(f"{label}{val}")
            rv.font.size = Pt(10)

    doc.add_paragraph()

    # Wykonawca
    p_w = doc.add_paragraph()
    run_w = p_w.add_run("WYKONAWCA:")
    run_w.bold = True
    run_w.font.size = Pt(11)

    for k, label in [("name", ""), ("address", ""), ("nip", "NIP: "), ("phone", "Tel: "), ("email", "E-mail: ")]:
        val = owner_data.get(k, "")
        if val:
            pv = doc.add_paragraph()
            pv.paragraph_format.space_before = Pt(2)
            pv.paragraph_format.space_after = Pt(2)
            rv = pv.add_run(f"{label}{val}")
            rv.font.size = Pt(10)

    doc.add_paragraph()

    # Wartości
    total = sum(_to_decimal(l.get("line_total_pln", 0)) for l in lines)
    kp = total * config.kp_percent / Decimal("100")
    zysk = (total + kp) * config.zysk_percent / Decimal("100")
    netto = total + kp + zysk
    vat = netto * config.vat_percent / Decimal("100")
    brutto = netto + vat

    val_table = doc.add_table(rows=0, cols=2)
    val_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    def _add_val_row(label: str, value: str, bold_val: bool = False):
        row = val_table.add_row()
        _set_cell_text(row.cells[0], label, bold=True, size=11)
        _set_cell_text(row.cells[1], f"{value} {config.currency}", bold=bold_val, size=11, align=WD_ALIGN_PARAGRAPH.RIGHT)

    _add_val_row("Wartość netto:", _format_pln(netto))
    _add_val_row(f"VAT ({config.vat_percent}%):", _format_pln(vat))
    _add_val_row("Wartość brutto:", _format_pln(brutto), bold_val=True)


# ---------------------------------------------------------------------------
# Estimate table
# ---------------------------------------------------------------------------

def _build_estimate_table(doc: Document, lines: List[Dict], cols: List[Tuple], config: DocxExportConfig):
    """Buduj tabelę kosztorysu z grupowaniem po rozdziałach."""
    # Header
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rh = p_header.add_run("TABELA KOSZTORYSU")
    rh.bold = True
    rh.font.size = Pt(12)

    # Grupowanie po rozdziałach
    chapters: Dict[str, List[Dict]] = {}
    for line in lines:
        prov = line.get("provenance", {}) or {}
        chapter = prov.get("chapter", "Pozycje bez przypisania")
        if chapter not in chapters:
            chapters[chapter] = []
        chapters[chapter].append(line)

    num_cols = len(cols)
    table = doc.add_table(rows=1, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_table_borders(table)

    # Set column widths
    for i, (_, width, _) in enumerate(cols):
        table.columns[i].width = Cm(width)

    # Header row
    hdr_row = table.rows[0]
    for i, (name, _, align) in enumerate(cols):
        cell = hdr_row.cells[i]
        _set_cell_text(cell, name, bold=True, align=ALIGN_MAP[align], size=9)
        _shade_cell(cell, "D4D4D8")

    # Data rows
    lp = 0
    row_idx = 0
    for chapter_name, chapter_lines in chapters.items():
        # Chapter header row
        ch_row = table.add_row()
        row_idx += 1
        # Merge all cells for chapter header
        ch_row.cells[0].merge(ch_row.cells[num_cols - 1])
        _set_cell_text(ch_row.cells[0], chapter_name, bold=True, size=10)
        _shade_cell(ch_row.cells[0], "E8E8EC")

        chapter_total = Decimal("0")

        for line in chapter_lines:
            lp += 1
            row_idx += 1
            data_row = table.add_row()

            prov = line.get("provenance", {}) or {}
            knr = prov.get("knr_code", "")
            desc = line.get("description", "")
            unit = line.get("unit", "")
            qty = _to_decimal(line.get("quantity", 0))
            uprice = _to_decimal(line.get("unit_price", 0))
            labor = _to_decimal(line.get("labor_pln", 0))
            material = _to_decimal(line.get("material_pln", 0))
            equipment = _to_decimal(line.get("equipment_pln", 0))
            total = _to_decimal(line.get("line_total_pln", 0))
            chapter_total += total

            # Fill cells based on template
            if config.template == "uproszczony":
                values = [
                    str(lp), knr, desc, unit,
                    _format_pln(qty), _format_pln(total),
                ]
            elif config.template == "szczegolowy":
                values = [
                    str(lp), knr, desc, unit,
                    _format_pln(qty), _format_pln(uprice),
                    _format_pln(labor), _format_pln(material),
                    _format_pln(equipment), _format_pln(total),
                ]
            else:  # ofertowy, inwestorski
                values = [
                    str(lp), knr, desc, unit,
                    _format_pln(qty), _format_pln(uprice),
                    _format_pln(labor), _format_pln(material),
                    _format_pln(equipment),
                ]

            for i, val in enumerate(values):
                cell = data_row.cells[i]
                _, _, align = cols[i]
                _set_cell_text(cell, val, align=ALIGN_MAP[align], size=9)

            # Alternating stripes
            if row_idx % 2 == 0:
                for i in range(num_cols):
                    _shade_cell(data_row.cells[i], "F4F4F5")

        # Subtotal row per chapter
        sub_row = table.add_row()
        row_idx += 1
        sub_row.cells[0].merge(sub_row.cells[num_cols - 2])
        _set_cell_text(sub_row.cells[0], f"Razem {chapter_name}:", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, size=9)
        _set_cell_text(sub_row.cells[num_cols - 1], _format_pln(chapter_total), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, size=9)
        _shade_cell(sub_row.cells[0], "E8E8EC")
        _shade_cell(sub_row.cells[num_cols - 1], "E8E8EC")


# ---------------------------------------------------------------------------
# Summary
# ---------------------------------------------------------------------------

def _build_summary(doc: Document, lines: List[Dict], config: DocxExportConfig):
    """Buduj sekcję podsumowania."""
    doc.add_paragraph()
    p_s = doc.add_paragraph()
    run_s = p_s.add_run("PODSUMOWANIE")
    run_s.bold = True
    run_s.font.size = Pt(12)

    total = sum(_to_decimal(l.get("line_total_pln", 0)) for l in lines)
    kp = total * config.kp_percent / Decimal("100")
    zysk = (total + kp) * config.zysk_percent / Decimal("100")
    netto = total + kp + zysk
    vat = netto * config.vat_percent / Decimal("100")
    brutto = netto + vat

    sum_table = doc.add_table(rows=0, cols=2)
    sum_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _apply_table_borders(sum_table)

    # Set widths
    sum_table.columns[0].width = Cm(10)
    sum_table.columns[1].width = Cm(5)

    def _add_sum_row(label: str, value: str, bold: bool = False, shade: Optional[str] = None):
        row = sum_table.add_row()
        _set_cell_text(row.cells[0], label, bold=bold, size=10)
        _set_cell_text(row.cells[1], f"{value} {config.currency}", bold=bold, align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)
        if shade:
            _shade_cell(row.cells[0], shade)
            _shade_cell(row.cells[1], shade)

    _add_sum_row("Wartość robót:", _format_pln(total))
    if config.kp_percent > 0:
        _add_sum_row(f"Koszty pośrednie ({config.kp_percent}%):", _format_pln(kp))
    if config.zysk_percent > 0:
        _add_sum_row(f"Zysk ({config.zysk_percent}%):", _format_pln(zysk))
    _add_sum_row("RAZEM NETTO:", _format_pln(netto), bold=True, shade="E8E8EC")
    _add_sum_row(f"VAT ({config.vat_percent}%):", _format_pln(vat))
    _add_sum_row("RAZEM BRUTTO:", _format_pln(brutto), bold=True, shade="D4D4D8")

    # Kwota słownie
    doc.add_paragraph()
    p_words = doc.add_paragraph()
    run_w = p_words.add_run("Słownie brutto: ")
    run_w.bold = True
    run_w.font.size = Pt(10)
    run_v = p_words.add_run(_amount_in_words(brutto))
    run_v.font.size = Pt(10)
    run_v.italic = True


# ---------------------------------------------------------------------------
# Merged elements table
# ---------------------------------------------------------------------------

def _build_merged_elements_table(doc: Document, lines: List[Dict], config: DocxExportConfig):
    """Tabela elementów scalonych — podsumowanie kosztów wg kategorii."""
    doc.add_paragraph()
    p_me = doc.add_paragraph()
    run_me = p_me.add_run("TABELA ELEMENTÓW SCALONYCH")
    run_me.bold = True
    run_me.font.size = Pt(12)

    # Grupowanie po rozdziałach
    chapters: Dict[str, Dict[str, Decimal]] = {}
    for line in lines:
        prov = line.get("provenance", {}) or {}
        chapter = prov.get("chapter", "Inne")
        if chapter not in chapters:
            chapters[chapter] = {"labor": Decimal("0"), "material": Decimal("0"), "equipment": Decimal("0"), "total": Decimal("0")}
        chapters[chapter]["labor"] += _to_decimal(line.get("labor_pln", 0))
        chapters[chapter]["material"] += _to_decimal(line.get("material_pln", 0))
        chapters[chapter]["equipment"] += _to_decimal(line.get("equipment_pln", 0))
        chapters[chapter]["total"] += _to_decimal(line.get("line_total_pln", 0))

    me_table = doc.add_table(rows=1, cols=5)
    me_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_table_borders(me_table)

    headers = ["Element", "Robocizna", "Materiały", "Sprzęt", "Razem"]
    for i, h in enumerate(headers):
        _set_cell_text(me_table.rows[0].cells[i], h, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade_cell(me_table.rows[0].cells[i], "D4D4D8")

    grand = {"labor": Decimal("0"), "material": Decimal("0"), "equipment": Decimal("0"), "total": Decimal("0")}
    for ch_name, vals in chapters.items():
        row = me_table.add_row()
        _set_cell_text(row.cells[0], ch_name, size=9)
        _set_cell_text(row.cells[1], _format_pln(vals["labor"]), align=WD_ALIGN_PARAGRAPH.RIGHT, size=9)
        _set_cell_text(row.cells[2], _format_pln(vals["material"]), align=WD_ALIGN_PARAGRAPH.RIGHT, size=9)
        _set_cell_text(row.cells[3], _format_pln(vals["equipment"]), align=WD_ALIGN_PARAGRAPH.RIGHT, size=9)
        _set_cell_text(row.cells[4], _format_pln(vals["total"]), align=WD_ALIGN_PARAGRAPH.RIGHT, size=9)
        for k in grand:
            grand[k] += vals[k]

    # Total row
    tot_row = me_table.add_row()
    _set_cell_text(tot_row.cells[0], "RAZEM:", bold=True, size=9)
    _set_cell_text(tot_row.cells[1], _format_pln(grand["labor"]), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, size=9)
    _set_cell_text(tot_row.cells[2], _format_pln(grand["material"]), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, size=9)
    _set_cell_text(tot_row.cells[3], _format_pln(grand["equipment"]), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, size=9)
    _set_cell_text(tot_row.cells[4], _format_pln(grand["total"]), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, size=9)
    for i in range(5):
        _shade_cell(tot_row.cells[i], "D4D4D8")


# ---------------------------------------------------------------------------
# Signatures
# ---------------------------------------------------------------------------

def _build_signatures(doc: Document, config: DocxExportConfig):
    """Buduj sekcję podpisów."""
    doc.add_paragraph()
    doc.add_paragraph()

    sig_table = doc.add_table(rows=2, cols=len(config.signatures))
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, sig in enumerate(config.signatures):
        # Label
        _set_cell_text(sig_table.rows[0].cells[i], sig + ":", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
        # Line for signature
        _set_cell_text(sig_table.rows[1].cells[i], "\n\n" + "." * 30, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

    # Date
    doc.add_paragraph()
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.LEFT
    est_date = config.estimate_date or date.today().strftime("%Y-%m-%d")
    rd = p_date.add_run(f"Data: {est_date}")
    rd.font.size = Pt(9)
